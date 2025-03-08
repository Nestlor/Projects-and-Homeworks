import os
os.chdir('d:/dev/python/lesson12')

#1)
import json
import requests

url = "https://jsonplaceholder.typicode.com/todos/"
response = requests.get(url)

if response.status_code == 200:
    data = response.json()
    with open("todos.json", "w", encoding="utf-8") as file:
        json.dump(data, file, indent=4)
    
    with open("todos.json", "r", encoding="utf-8") as file:
        loaded_data = json.load(file)
        print(loaded_data)

# 2)
from docx import Document
from docx.shared import Pt

doc = Document()
para = doc.add_paragraph()
run = para.add_run("Hello World")
run.bold = True

doc.save("hello.docx")

doc = Document("hello.docx")

bold_text = ""
for para in doc.paragraphs:
    for run in para.runs:
        if run.bold:
            bold_text += run.text + " "

print("Жирный текст:", bold_text.strip())

doc = Document()
para = doc.add_paragraph("текст с новым шрифтом и размером")
run = para.runs[0]
run.font.name = "Arial"
run.font.size = Pt(14)

doc.save("new_file.docx")